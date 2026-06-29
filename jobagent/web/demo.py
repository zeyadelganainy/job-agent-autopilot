"""Demo mode: a self-contained, no-live-calls sandbox for recruiters to click through.

A demo session is backed by its own SQLite DB (`DEMO_DB`) — separate from the real
`jobs.db` — seeded with realistic sample data, plus premade documents written under
`<output>/demo`. Every "live" action (scan / generate) is *simulated* here: no LLM,
SMTP, or network call is ever made. The first simulated run adds 2 new matches and
prepares documents; pressing it again returns a friendly "this is a demo" notice.

The whole sandbox is reset and reseeded each time someone clicks "View demo", so every
recruiter starts from the same clean slate.
"""
import json
import time
from datetime import datetime, timedelta, timezone
from pathlib import Path

from ..config import ROOT
from ..store import Store

# Each demo visitor gets their OWN freshly-seeded database under DEMO_DIR, so no one can
# alter the demo for anyone else. Stale session DBs are pruned on each new entry.
DEMO_DIR = ROOT / ".demo"
DEMO_NOTICE = ("Demo mode — this is a private sandbox just for you. It's a simulation, "
               "no live API calls are made. Click around freely!")
SESSION_MAX_AGE = 6 * 3600          # prune demo DBs older than this


def is_demo(sess) -> bool:
    return bool(sess) and sess.get("role") == "demo"


def session_db(demo_id: str) -> str:
    """Absolute path to a demo session's SQLite DB (Store treats an absolute path as-is)."""
    return str((DEMO_DIR / f"demo_{demo_id or 'shared'}.db").resolve())


def _wipe(path: str):
    for suffix in ("", "-wal", "-shm"):
        p = Path(path + suffix)
        if p.exists():
            try:
                p.unlink()
            except OSError:
                pass


def prune(max_age: float = SESSION_MAX_AGE):
    """Delete stale demo session DBs so old sessions can't pile up (or be revisited)."""
    if not DEMO_DIR.exists():
        return
    now = time.time()
    for p in DEMO_DIR.glob("demo_*.db"):
        try:
            if now - p.stat().st_mtime > max_age:
                _wipe(str(p))
        except OSError:
            pass


def ensure_session(output_dir: Path, demo_id: str) -> str:
    """Return a ready demo DB for this session, seeding a fresh one if it's missing
    (e.g. after a prune) — so a returning demo cookie always lands on clean data."""
    path = session_db(demo_id)
    if not Path(path).exists():
        reset_and_seed(output_dir, demo_id)
    return path


def _recent(days_ago: int = 1) -> str:
    return (datetime.now(timezone.utc) - timedelta(days=days_ago)).isoformat()


def _insert_job(conn, job_id, title, company, score, status, *, location, url,
                posted=None, reasons="", gaps="", docs=None, flag_reason=None,
                applied_at=None):
    conn.execute(
        """INSERT OR REPLACE INTO jobs
               (job_id, source, title, company, url, location, description, posted,
                score, reasons, gaps, status, docs, applied_at, flag_reason, created_at)
           VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,CURRENT_TIMESTAMP)""",
        (job_id, "demo", title, company, url, location,
         f"{title} at {company}. (Sample posting used for the interactive demo.)",
         posted or _recent(), score, reasons, gaps,
         status, json.dumps(docs) if docs else None, applied_at, flag_reason))


# ---- premade documents -------------------------------------------------------
def _write_docx(path: Path, heading: str, paragraphs: list[str]):
    from docx import Document
    doc = Document()
    doc.add_heading(heading, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))


def _doc_paths(output_dir: Path, title: str, company: str) -> list[str]:
    """Ensure a premade résumé + cover letter exist for this role; return their paths."""
    base = Path(output_dir) / "demo"
    base.mkdir(parents=True, exist_ok=True)
    safe = "".join(c if c.isalnum() else "_" for c in f"{company}_{title}").strip("_")[:60]
    resume = base / f"Resume_{safe}.docx"
    cover = base / f"CoverLetter_{safe}.docx"
    if not resume.exists():
        _write_docx(resume, f"{title} — Résumé", [
            "Alex Rivera · alex.rivera@example.com · Vancouver, BC",
            "",
            "SUMMARY",
            f"Software engineer tailored for the {title} role at {company}. This is a "
            "sample document generated for the interactive demo — in the live app it is a "
            "fully tailored résumé built from your master profile.",
            "",
            "EXPERIENCE",
            "Software Engineer — Example Corp (2022–present): shipped backend services, "
            "improved reliability, mentored juniors.",
            "Developer — Sample Labs (2020–2022): built APIs and internal tooling.",
            "",
            "SKILLS: Python, TypeScript, PostgreSQL, AWS, Docker.",
        ])
    if not cover.exists():
        _write_docx(cover, f"Cover Letter — {company}", [
            f"Dear {company} Hiring Team,",
            "",
            f"I'm excited to apply for the {title} role. (Sample cover letter generated for "
            "the demo — the live app writes this in your own voice from your writing samples.)",
            "",
            "I bring hands-on experience building reliable backend systems and a track record "
            "of shipping. I'd love to contribute to your team.",
            "",
            "Warm regards,",
            "Alex Rivera",
        ])
    return [str(resume.resolve()), str(cover.resolve())]


# ---- seeding -----------------------------------------------------------------
def _seed(store: Store, output_dir: Path):
    conn = store.conn
    # 3 jobs on the board (status 'sent'). "Run agent now" adds 2 more (see run_scan).
    _insert_job(conn, "demo-1", "Backend Engineer", "Stripe", 88, "sent",
                location="Remote · Canada",
                url="https://job-boards.greenhouse.io/stripe/jobs/demo",
                reasons="Strong match on Python/distributed systems and remote-Canada.")
    _insert_job(conn, "demo-2", "Software Engineer, New Grad", "Shopify", 82, "sent",
                location="Toronto, ON",
                url="https://job-boards.greenhouse.io/shopify/jobs/demo",
                reasons="New-grad friendly; your stack and location line up well.")
    _insert_job(conn, "demo-3", "Full-Stack Developer", "Wealthsimple", 76, "sent",
                location="Toronto, ON",
                url="https://job-boards.greenhouse.io/wealthsimple/jobs/demo",
                reasons="Good overlap, slightly more frontend than your core.")

    # A triage item, so the "Needs your attention" queue isn't empty. Docs are prepared
    # but the posting lives on an external portal, so you apply manually.
    _insert_job(conn, "demo-att", "Platform Engineer", "Cloudgrid", 85, "needs_attention",
                location="Remote",
                url="https://cloudgrid.wd1.myworkdayjobs.com/careers/demo",
                reasons="Great match on infra/platform experience.",
                flag_reason="Manual application portal (Workday/external) — apply yourself",
                docs=_doc_paths(output_dir, "Platform Engineer", "Cloudgrid"))

    # Tracker history — a realistic funnel spread across several months so the Insights
    # charts (applications over time + by stage) have something to show.
    today = datetime.now()
    history = [
        # (role, company, days_ago, stage, location)
        ("Software Engineer", "Asana", 2, "Applied", "Vancouver, BC"),
        ("Backend Engineer", "Datadog", 5, "Applied", "Remote"),
        ("Platform Engineer", "Cloudflare", 9, "Screening", "Remote · Canada"),
        ("Backend Engineer", "Notion", 14, "Interview", "Remote"),
        ("Full-Stack Engineer", "Linear", 18, "Applied", "Remote"),
        ("Software Engineer II", "Twilio", 24, "Rejected", "Toronto, ON"),
        ("Backend Engineer", "GitLab", 31, "Interview", "Remote"),
        ("Software Engineer", "Vercel", 38, "Offer", "Remote"),
        ("New Grad SWE", "Shopify", 47, "Applied", "Ottawa, ON"),
        ("Backend Developer", "Ramp", 58, "Ghosted", "Remote"),
        ("Platform Engineer", "Databricks", 69, "Rejected", "Vancouver, BC"),
        ("Software Engineer", "Atlassian", 82, "Interview", "Remote · Canada"),
        ("Backend Engineer", "Plaid", 96, "Applied", "Toronto, ON"),
        ("Full-Stack Engineer", "Brex", 118, "Ghosted", "Remote"),
        ("Software Engineer", "Snowflake", 142, "Rejected", "Remote"),
    ]
    for role, company, days_ago, stage, location in history:
        store.add_application({
            "role": role, "company": company,
            "applied_date": (today - timedelta(days=days_ago)).strftime("%Y-%m-%d"),
            "stage": stage, "location": location, "notes": "",
            "url": f"https://example.com/{company.lower().replace(' ', '')}"})
    conn.commit()


def reset_and_seed(output_dir: Path, demo_id: str):
    """Wipe and reseed one demo session's DB so this visitor starts from a clean slate."""
    DEMO_DIR.mkdir(parents=True, exist_ok=True)
    path = session_db(demo_id)
    _wipe(path)
    store = Store(path)
    try:
        _seed(store, output_dir)
    finally:
        store.close()


# ---- simulated "live" actions ------------------------------------------------
def run_scan(store: Store, output_dir: Path) -> str:
    """Simulate one agent run. First call adds 2 matches + prepares 2 doc sets; after
    that, just returns the demo notice (no live calls — ever)."""
    if store.list_runs(1):                       # already simulated once this session
        return DEMO_NOTICE
    run_id = store.start_run()
    new = [("demo-4", "Backend Developer", "Figma", 91, "Remote",
            "https://job-boards.greenhouse.io/figma/jobs/demo"),
           ("demo-5", "Software Engineer II", "1Password", 84, "Remote · Canada",
            "https://jobs.ashbyhq.com/1password/demo")]
    for jid, title, company, score, loc, url in new:
        _insert_job(store.conn, jid, title, company, score, "sent",
                    location=loc, url=url,
                    reasons="Strong match on your core stack and location.")
    store.conn.commit()

    # Prepare documents for the two highest-scoring board jobs that don't have any yet.
    sent = sorted((r for r in store.by_status("sent") if not r["docs"]),
                  key=lambda r: r["score"] or 0, reverse=True)
    prepared = 0
    for r in sent[:2]:
        store.record_docs(r["job_id"], _doc_paths(output_dir, r["title"], r["company"]))
        prepared += 1

    store.finish_run(run_id, status="ok", scanned=5, scored=2, matched=2, generated=prepared)
    return (f"Demo run complete — added 2 new matches and prepared {prepared} document "
            "sets. No live calls were made.")


def simulate_generate(store: Store, output_dir: Path, job_ids: list[str]) -> str:
    """Attach premade documents to the selected jobs (no LLM call)."""
    n = 0
    for jid in job_ids:
        row = store.get(jid)
        if not row:
            continue
        store.record_docs(jid, _doc_paths(output_dir, row["title"], row["company"]))
        n += 1
    return (f"Simulated document generation for {n} role(s) — no live calls were made. "
            "Find the files under “Ready to apply” and in Docs.")


def simulate_adhoc(store: Store, output_dir: Path, title: str, company: str) -> dict:
    """Simulate the ad-hoc Generate page: premade docs for a synthetic job."""
    title = (title or "Pasted role").strip()
    company = (company or "Pasted company").strip()
    jid = "demo-adhoc-" + str(int(time.time()))
    _insert_job(store.conn, jid, title, company, 80, "generated",
                location="", url="", reasons="Generated in demo mode.")
    store.conn.commit()
    paths = _doc_paths(output_dir, title, company)
    store.record_docs(jid, paths)
    return {"ok": True, "demo": True, "job": dict(store.get(jid)),
            "docs": [{"name": Path(p).name, "path": p} for p in paths]}
