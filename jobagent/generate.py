"""Generate a tailored cover letter + resume for a picked job, in the user's voice.

Content comes from the master résumé (profile/master.md, the source of truth). The
resume is produced as structured JSON, then rendered into a .docx that matches the
user's own template (profile/resume.docx) — its fonts, margins, and two-column layout.

Outputs to output/<job_id>/, named [date]_[title]_[company]_[resume|coverLetter]:
  e.g. 20260625_software_developer_d2l_resume.docx  and  ..._coverLetter.docx
"""
import re
import warnings
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from .config import ROOT
from .llm import chat, extract_json

# ---------------------------------------------------------------- cover letter
COVER_SYSTEM = (
    "You write cover letters that sound like the candidate, not like an AI. "
    "Match the tone, rhythm, and vocabulary of the writing samples provided. "
    "No clichés ('I am thrilled', 'fast-paced environment'), no invented facts — "
    "use only what's in the master résumé below. 3 short paragraphs, ~250 words. "
    "Return only the letter."
)

COVER_TEMPLATE = """Writing samples that show the candidate's voice:
{samples}

Candidate master résumé (source of truth — use only facts from here):
{master}

Job: {title} at {company}
Job description:
{description}

Write a tailored cover letter for this job, in the candidate's voice."""

# --------------------------------------------------------------------- resume
RESUME_SYSTEM = (
    "You tailor a résumé to a specific job using ONLY the candidate's master résumé "
    "as the source of truth. Obey the operating instructions embedded in the master "
    "file: never fabricate; select, reorder, and lightly rephrase real content; "
    "reproduce every number, date, and name exactly; prefer quantified bullets; if the "
    "job needs something not in the master, omit it rather than invent it. "
    "It MUST fit on ONE page — be selective, not exhaustive. "
    "Output ONLY a JSON object — no prose, no code fences."
)

RESUME_TEMPLATE = """Master résumé (source of truth):
{master}

Target job: {title} at {company}
Job description:
{description}

Produce a tailored ONE-PAGE résumé as JSON with EXACTLY this shape:
{{
  "summary": "2-3 line professional summary tailored to this job",
  "skills": [{{"label": "Languages", "items": "comma-separated values"}}],
  "experience": [{{"role": "", "org": "", "location": "", "dates": "", "bullets": [""]}}],
  "projects": [{{"name": "", "tech": "", "dates": "", "link": "", "bullets": [""]}}],
  "education": [{{"school": "", "degree": "", "location": "", "dates": "", "bullets": [""]}}]
}}

ONE PAGE IS A HARD LIMIT. To stay within it:
- At most 3-4 bullets per experience entry; keep the most relevant, quantified ones.
- At most 2-3 projects, the most relevant to THIS job; <= 3 bullets each.
- One concise line per skill category; drop categories irrelevant to the role.
- Keep the summary to 2-3 lines.

"link" MUST be a URL that appears VERBATIM in the master file (e.g. a project's GitHub
link); if the master has no link for that project, use "". Never invent or guess a URL.
"dates" is a single display string (e.g. "July 2023 – June 2024"). Use only facts from
the master; omit any field you have no real content for."""

# Right edge of the text column on US Letter with the template's 0.5" margins.
_RIGHT_TAB = Inches(7.5)


# -------------------------------------------------------- low-level docx helpers
def _clear_body(doc: Document) -> None:
    """Remove the template's own paragraphs/tables but keep section (margin) props.

    Also drop the template's hyperlink relationships — clearing the body leaves them
    dangling in the part, which would embed stray (and possibly wrong) URLs.
    """
    body = doc.element.body
    for child in list(body):
        if child.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(child)
    for rid, rel in list(doc.part.rels.items()):
        if rel.reltype == RT.HYPERLINK:
            doc.part.drop_rel(rid)


def _par(doc: Document) -> Paragraph:
    """Append a paragraph *before* the trailing sectPr so margins/page survive."""
    body = doc.element.body
    p = OxmlElement("w:p")
    sect = body.find(qn("w:sectPr"))
    if sect is not None:
        sect.addprevious(p)
    else:
        body.append(p)
    para = Paragraph(p, doc._body)
    para.paragraph_format.space_after = Pt(2)
    return para


def _run(para: Paragraph, text: str, *, bold=False, italic=False, size=10):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return r


def _url(u: str) -> str:
    """Add a scheme so the URL is clickable (master stores bare 'github.com/...')."""
    u = str(u or "").strip()
    if not u or u.startswith(("http://", "https://", "mailto:")):
        return u
    return ("mailto:" + u) if "@" in u and "/" not in u else ("https://" + u)


def _hr(para: Paragraph, size="6"):
    """Draw a horizontal rule under `para` (a paragraph bottom border)."""
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)        # eighths of a point (6 = 0.75pt)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "808080")
    pbdr.append(bottom)
    # pBdr must precede spacing/jc/ind in the CT_PPr child order, so insert first.
    para._p.get_or_add_pPr().insert(0, pbdr)


def _hyperlink(para: Paragraph, url: str, text: str, *, bold=False, italic=False, size=10):
    """Append a real (clickable) hyperlink run to `para`."""
    r_id = para.part.relate_to(_url(url), RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rPr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if size:
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    link.append(run)
    para._p.append(link)
    return link


def _two_col(doc, left, right, *, bold_left=False, italic=False, size=10):
    """A line with `left` text and `right` text pushed to the right margin via a tab."""
    p = _par(doc)
    p.paragraph_format.tab_stops.add_tab_stop(_RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    _run(p, left, bold=bold_left, italic=italic, size=size)
    if right:
        _run(p, "\t" + right, italic=italic, size=size)
    return p


def _section(doc, title: str):
    p = _par(doc)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    _run(p, title.upper(), bold=True, size=11)
    _hr(p)                       # horizontal rule under every section header


def _bullet(doc, text: str):
    p = _par(doc)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.15)
    pf.space_after = Pt(1)
    _run(p, "•  " + text, size=10)


# ------------------------------------------------------------- resume rendering
def _contact_line(identity: dict) -> str:
    parts = [identity.get(k) for k in ("location", "phone", "email", "github", "linkedin")]
    return "  |  ".join(p for p in parts if p)


def _render_contact(doc, identity: dict):
    """Centered contact paragraph; email/github/linkedin/portfolio are clickable."""
    # (key, is_link)
    fields = [("location", False), ("phone", False), ("email", True),
              ("github", True), ("linkedin", True), ("portfolio", True)]
    parts = [(identity.get(k), is_link) for k, is_link in fields if identity.get(k)]
    if not parts:
        return None
    p = _par(doc)
    p.alignment = 1  # center
    for i, (val, is_link) in enumerate(parts):
        if i:
            _run(p, "  |  ", size=10)
        if is_link:
            _hyperlink(p, str(val), str(val), size=10)
        else:
            _run(p, str(val), size=10)
    return p


def _norm_skills(skills):
    """Accept either a list of {label, items} or a {label: items} dict."""
    if isinstance(skills, dict):
        return [{"label": k, "items": v} for k, v in skills.items()]
    return skills or []


def _strip_fabricated_links(data: dict, master: str) -> None:
    """Drop any project link not present verbatim in the master (anti-fabrication)."""
    m = (master or "").lower()
    for pr in data.get("projects") or []:
        link = str(pr.get("link") or "").strip()
        core = link.lower().split("://", 1)[-1].rstrip("/")
        if core and core not in m:
            pr["link"] = ""


def _render_resume_docx(data: dict, identity: dict, template: Path, out_path: Path):
    doc = Document(str(template)) if template.exists() else Document()
    _clear_body(doc)

    # Tighten default spacing so the résumé stays on one page.
    with warnings.catch_warnings():            # template's default style id is "Normal"
        warnings.simplefilter("ignore")
        nf = doc.styles["Normal"].paragraph_format
    nf.space_before = Pt(0)
    nf.space_after = Pt(2)
    nf.line_spacing = 1.0

    # Header: centered name + hyperlinked contact, with a rule beneath it.
    name_p = _par(doc)
    name_p.alignment = 1  # center
    _run(name_p, identity.get("name", ""), bold=True, size=18)
    _hr(_render_contact(doc, identity) or name_p)

    if data.get("summary"):
        _par(doc).add_run(str(data["summary"])).font.size = Pt(10)

    skills = _norm_skills(data.get("skills"))
    if skills:
        _section(doc, "Skills & Technologies")
        for cat in skills:
            label, items = cat.get("label"), cat.get("items")
            if not items:
                continue
            p = _par(doc)
            if label:
                _run(p, f"{label}: ", bold=True, size=10)
            _run(p, str(items), size=10)

    if data.get("experience"):
        _section(doc, "Experience")
        for e in data["experience"]:
            _two_col(doc, e.get("role", ""), e.get("location", ""), bold_left=True)
            sub = e.get("org", "")
            if sub or e.get("dates"):
                _two_col(doc, sub, e.get("dates", ""), italic=True)
            for b in e.get("bullets", []):
                _bullet(doc, b)

    if data.get("projects"):
        _section(doc, "Projects")
        for pr in data["projects"]:
            p = _par(doc)
            p.paragraph_format.tab_stops.add_tab_stop(_RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
            if pr.get("link"):
                _hyperlink(p, pr["link"], pr.get("name", ""), bold=True, size=10)
            else:
                _run(p, pr.get("name", ""), bold=True, size=10)
            if pr.get("tech"):
                _run(p, " | " + pr["tech"], italic=True, size=10)
            if pr.get("dates"):
                _run(p, "\t" + pr["dates"], italic=True, size=10)
            for b in pr.get("bullets", []):
                _bullet(doc, b)

    if data.get("education"):
        _section(doc, "Education")
        for ed in data["education"]:
            _two_col(doc, ed.get("school", ""), ed.get("location", ""), bold_left=True)
            deg = ed.get("degree", "")
            if deg or ed.get("dates"):
                _two_col(doc, deg, ed.get("dates", ""), italic=True)
            for b in ed.get("bullets", []):
                _bullet(doc, b)

    doc.save(str(out_path))


def _render_cover_docx(text: str, identity: dict, template: Path, out_path: Path):
    """Render the cover letter prose into a .docx with the same header/template look."""
    doc = Document(str(template)) if template.exists() else Document()
    _clear_body(doc)
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        nf = doc.styles["Normal"].paragraph_format
    nf.space_before = Pt(0)
    nf.space_after = Pt(6)
    nf.line_spacing = 1.15

    name_p = _par(doc)
    name_p.alignment = 1  # center
    _run(name_p, identity.get("name", ""), bold=True, size=16)
    _hr(_render_contact(doc, identity) or name_p)
    _par(doc)  # blank spacer line

    for block in re.split(r"\n\s*\n", (text or "").strip()):
        block = " ".join(block.split())   # collapse internal line breaks
        if block:
            p = _par(doc)
            p.paragraph_format.space_after = Pt(6)
            _run(p, block, size=10.5)
    doc.save(str(out_path))


def _resume_to_md(data: dict, identity: dict) -> str:
    """A plain-Markdown mirror of the structured resume, for quick reading/diffing."""
    out = [f"# {identity.get('name', '')}", _contact_line(identity), ""]
    if data.get("summary"):
        out += [str(data["summary"]), ""]
    for cat in _norm_skills(data.get("skills")):
        if cat.get("items"):
            out.append(f"**{cat.get('label', 'Skills')}:** {cat['items']}")
    if data.get("skills"):
        out.append("")
    for header, key, title_key, sub_key in (
        ("Experience", "experience", "role", "org"),
        ("Projects", "projects", "name", "tech"),
        ("Education", "education", "school", "degree"),
    ):
        items = data.get(key) or []
        if not items:
            continue
        out.append(f"## {header}")
        for it in items:
            title = it.get(title_key, "")
            dates = it.get("dates", "")
            out.append(f"### {title}{(' — ' + dates) if dates else ''}")
            sub = it.get(sub_key)
            if sub:
                out.append(f"*{sub}*")
            if it.get("link"):
                out.append(f"[{it['link']}]({_url(it['link'])})")
            for b in it.get("bullets", []):
                out.append(f"- {b}")
            out.append("")
    return "\n".join(out).strip() + "\n"


def _slug(text, maxlen: int = 40) -> str:
    """lowercase, non-alphanumeric -> underscore, collapsed and trimmed."""
    s = re.sub(r"[^a-z0-9]+", "_", str(text or "").lower()).strip("_")
    return s[:maxlen].strip("_") or "x"


def _doc_basename(job_row) -> str:
    """[date]_[title]_[company], e.g. 20260625_software_developer_d2l."""
    return f"{date.today():%Y%m%d}_{_slug(job_row['title'])}_{_slug(job_row['company'])}"


# ------------------------------------------------------------------- entry point
def generate(job_row, profile: dict, samples: str, master: str, models: dict,
             paths: dict) -> list[str]:
    job_id = job_row["job_id"]
    folder = ROOT / paths["output"] / job_id
    folder.mkdir(parents=True, exist_ok=True)
    desc = (job_row["description"] or "")[:6000]
    identity = profile.get("identity", profile)   # tolerate flat or nested profile
    template = ROOT / paths["template"]
    base = _doc_basename(job_row)                  # [date]_[title]_[company]
    out_paths = []

    # Cover letter -> .docx (falls back to .md only if rendering fails).
    cover = chat(COVER_SYSTEM, COVER_TEMPLATE.format(
        samples=samples or "(no samples provided)", master=master,
        title=job_row["title"], company=job_row["company"], description=desc), models)
    try:
        cover_path = folder / f"{base}_coverLetter.docx"
        _render_cover_docx(cover, identity, template, cover_path)
    except Exception as e:
        print(f"[generate] cover docx render failed: {e}; saving markdown")
        cover_path = folder / f"{base}_coverLetter.md"
        cover_path.write_text(cover, encoding="utf-8")
    out_paths.append(str(cover_path))

    # Resume: structured JSON -> template-matching .docx.
    raw = chat(RESUME_SYSTEM, RESUME_TEMPLATE.format(
        master=master, title=job_row["title"],
        company=job_row["company"], description=desc), models)
    try:
        data = extract_json(raw)
    except Exception as e:
        print(f"[generate] resume JSON parse failed ({e}); saving raw output.")
        md = folder / f"{base}_resume.md"
        md.write_text(raw, encoding="utf-8")
        return out_paths + [str(md)]

    _strip_fabricated_links(data, master)   # never let invented URLs onto the résumé
    try:
        resume_path = folder / f"{base}_resume.docx"
        _render_resume_docx(data, identity, template, resume_path)
    except Exception as e:
        print(f"[generate] resume docx render failed: {e}; saving markdown")
        resume_path = folder / f"{base}_resume.md"
        resume_path.write_text(_resume_to_md(data, identity), encoding="utf-8")
    out_paths.append(str(resume_path))

    return out_paths
