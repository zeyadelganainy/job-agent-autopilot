import json
from datetime import date
from pathlib import Path

import docx

import jobagent.generate as g


def test_norm_skills_accepts_dict_or_list():
    assert g._norm_skills({"Languages": "a, b"}) == [{"label": "Languages", "items": "a, b"}]
    lst = [{"label": "L", "items": "x"}]
    assert g._norm_skills(lst) == lst
    assert g._norm_skills(None) == []


def test_contact_line_skips_missing_fields():
    line = g._contact_line({"location": "Kelowna", "email": "e", "github": "gh"})
    assert line == "Kelowna  |  e  |  gh"


def test_contact_line_includes_website():
    line = g._contact_line({"email": "e", "website": "zeyad.ca"})
    assert line == "e  |  zeyad.ca"


def test_render_contact_renders_website_as_link(tmp_path):
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    d = Document()
    g._clear_body(d)
    g._render_contact(d, {"email": "e@x.com", "website": "zeyad.ca"})
    links = [r._target for r in d.part.rels.values() if r.reltype == RT.HYPERLINK]
    assert any("zeyad.ca" in t for t in links)


def test_resume_to_md_structure():
    data = {"summary": "S", "skills": [{"label": "L", "items": "x"}],
            "experience": [{"role": "R", "org": "O", "dates": "D", "bullets": ["b1"]}]}
    md = g._resume_to_md(data, {"name": "Zeyad"})
    assert "# Zeyad" in md and "## Experience" in md and "- b1" in md


def test_strip_fabricated_links_keeps_only_master_urls():
    master = "Project X\nLink: github.com/me/realproj\nmore text"
    data = {"projects": [
        {"name": "Real", "link": "https://github.com/me/realproj"},   # in master -> kept
        {"name": "Fake", "link": "https://wealthwise.vercel.app"},    # not in master -> dropped
        {"name": "None", "link": ""},
    ]}
    g._strip_fabricated_links(data, master)
    assert data["projects"][0]["link"] == "https://github.com/me/realproj"
    assert data["projects"][1]["link"] == ""
    assert data["projects"][2]["link"] == ""


def test_clear_body_drops_orphan_hyperlinks():
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    d = Document()
    g._hyperlink(d.add_paragraph(), "https://example.com", "x")
    assert any(r.reltype == RT.HYPERLINK for r in d.part.rels.values())
    g._clear_body(d)
    assert not any(r.reltype == RT.HYPERLINK for r in d.part.rels.values())


def test_url_normalization():
    assert g._url("github.com/x") == "https://github.com/x"
    assert g._url("https://x") == "https://x"
    assert g._url("a@b.com") == "mailto:a@b.com"
    assert g._url("") == ""


def test_render_docx_without_template_is_valid(tmp_path):
    data = {"summary": "S", "skills": [{"label": "L", "items": "x"}],
            "experience": [{"role": "R", "org": "O", "location": "loc",
                            "dates": "D", "bullets": ["b1"]}]}
    out = tmp_path / "r.docx"
    g._render_resume_docx(data, {"name": "Zeyad", "email": "e"},
                          Path("nonexistent.docx"), out)
    d = docx.Document(str(out))
    assert [c.tag.split('}')[1] for c in d.element.body][-1] == "sectPr"
    text = "\n".join(p.text for p in d.paragraphs)
    assert "Zeyad" in text and "EXPERIENCE" in text and "b1" in text


def test_render_docx_has_rules_and_project_hyperlink(tmp_path):
    data = {"summary": "S", "skills": [{"label": "L", "items": "x"}],
            "projects": [{"name": "Cool Project", "tech": "Python",
                          "link": "github.com/me/proj", "bullets": ["did a thing"]}]}
    out = tmp_path / "r.docx"
    g._render_resume_docx(data, {"name": "Z", "github": "github.com/me"},
                          Path("nonexistent.docx"), out)
    xml = docx.Document(str(out)).element.xml.lower()
    assert "w:pbdr" in xml          # horizontal rules (header + section headers)
    assert "hyperlink" in xml       # clickable project + contact links


def test_slug_and_basename():
    assert g._slug("Software Developer - New Graduate") == "software_developer_new_graduate"
    assert g._slug("D2L") == "d2l"
    assert g._slug("") == "x"
    base = g._doc_basename({"title": "Backend Engineer", "company": "D2L"})
    assert base == date.today().strftime("%Y%m%d") + "_backend_engineer_d2l"


def test_generate_writes_resume_and_cover_docx(tmp_path, monkeypatch):
    resume_json = json.dumps({
        "summary": "S", "skills": [{"label": "L", "items": "x"}],
        "experience": [{"role": "R", "org": "O", "dates": "D", "bullets": ["b"]}],
    })
    seq = iter(["Dear team, I am excited.\n\nSecond paragraph.", resume_json])
    monkeypatch.setattr(g, "chat", lambda s, u, m: next(seq))

    job = {"job_id": "abc", "title": "Backend Engineer", "company": "D2L", "description": "d"}
    paths = g.generate(job, {"identity": {"name": "Z", "email": "e"}}, "", "master",
                       {"claude": "c", "gemini": "x"},
                       {"output": str(tmp_path), "template": "nonexistent.docx"})
    names = sorted(Path(p).name for p in paths)
    prefix = date.today().strftime("%Y%m%d") + "_backend_engineer_d2l_"
    assert len(names) == 2
    assert names[0] == prefix + "coverLetter.docx"
    assert names[1] == prefix + "resume.docx"
    cover = next(p for p in paths if p.endswith("_coverLetter.docx"))
    text = "\n".join(p.text for p in docx.Document(cover).paragraphs)
    assert "Dear team" in text and "Second paragraph" in text


def test_generate_falls_back_when_resume_json_unparseable(tmp_path, monkeypatch):
    seq = iter(["cover text", "not json at all"])
    monkeypatch.setattr(g, "chat", lambda s, u, m: next(seq))
    job = {"job_id": "z", "title": "T", "company": "C", "description": "d"}
    paths = g.generate(job, {"identity": {"name": "Z"}}, "", "m",
                       {"claude": "c", "gemini": "x"},
                       {"output": str(tmp_path), "template": "nonexistent.docx"})
    names = {Path(p).name for p in paths}
    assert any(n.endswith("_coverLetter.docx") for n in names)   # cover -> docx
    assert any(n.endswith("_resume.md") for n in names)          # resume falls back to md
    assert not any(n.endswith("_resume.docx") for n in names)
